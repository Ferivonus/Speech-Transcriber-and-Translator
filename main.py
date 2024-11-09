import speech_recognition as sr
from docx import Document
from datetime import datetime
from googletrans import Translator
import tkinter as tk
from tkinter import scrolledtext, messagebox
import threading

def initialize_document():
    """Initializes and returns a new Word document with a header."""
    document = Document()
    document.add_heading('Speech Recognition Transcription', level=1)
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    document.add_paragraph("Transcriptions (Turkish - English):\n", style="Heading 2")
    return document

def transcribe_and_translate(recognizer, document, translator, output_text, stop_event):
    """Listens to audio, transcribes, translates, and displays in the GUI."""
    with sr.Microphone() as source:
        while not stop_event.is_set():  # Loop will continue until stop_event is set
            try:
                print("Listening...")
                audio = recognizer.listen(source)
                text_tr = recognizer.recognize_google(audio, language='tr-TR')
                print("You said (TR): " + text_tr)

                # Translate to English
                text_en = translator.translate(text_tr, src='tr', dest='en').text
                print("Translation (EN): " + text_en)

                # Save both Turkish and English transcription with timestamp
                timestamp = datetime.now().strftime('%H:%M:%S')
                document.add_paragraph(f"[{timestamp}] TR: {text_tr}")
                document.add_paragraph(f"[{timestamp}] EN: {text_en}")

                # Display in GUI
                output_text.insert(tk.END, f"[{timestamp}] TR: {text_tr}\n")
                output_text.insert(tk.END, f"[{timestamp}] EN: {text_en}\n\n")
                output_text.yview(tk.END)

            except sr.UnknownValueError:
                messagebox.showinfo("Error", "Could not understand the audio. Please try again.")
            except sr.RequestError as e:
                messagebox.showinfo("Error", f"Could not request results; {e}")

def save_transcriptions(document):
    """Save the document to a .docx file."""
    filename = f"transcriptions_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    document.save(filename)
    messagebox.showinfo("Saved", f"All transcriptions have been saved to '{filename}'.")

def start_transcription(stop_event):
    """Callback to start the transcription process."""
    stop_event.clear()  # Clear the stop event to start the transcription
    threading.Thread(target=transcribe_and_translate, args=(recognizer, document, translator, output_text, stop_event), daemon=True).start()

def stop_transcription(stop_event):
    """Callback to stop the transcription process."""
    stop_event.set()  # Set the stop event to stop the transcription loop

# Initialize recognizer, document, and translator
recognizer = sr.Recognizer()
document = initialize_document()
translator = Translator()

# Initialize GUI
root = tk.Tk()
root.title("Bilingual Speech Transcriber")
root.geometry("600x400")

# Add a text area for displaying transcriptions
output_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=15, font=("Arial", 10))
output_text.pack(pady=10)

# Add buttons to control transcription
start_button = tk.Button(root, text="Start Transcription", command=lambda: start_transcription(stop_event), font=("Arial", 12),
                         bg="lightgreen")
start_button.pack(pady=5)

stop_button = tk.Button(root, text="Stop Transcription", command=lambda: stop_transcription(stop_event), font=("Arial", 12),
                        bg="lightcoral")
stop_button.pack(pady=5)

save_button = tk.Button(root, text="Save Transcriptions", command=lambda: save_transcriptions(document),
                        font=("Arial", 12), bg="lightblue")
save_button.pack(pady=5)

# Initialize stop_event to control the transcription loop
stop_event = threading.Event()

# Run the GUI
root.mainloop()